import io
import pytest
from pathlib import Path
from unittest.mock import patch
from fastapi.testclient import TestClient

from backend.app.main import app
from backend.app.services import text_extraction_service
import backend.app.services.storage_service as storage_service

client = TestClient(app)


# ---------------------------------------------------------------------------
# Fixtures — real file content for each supported type
# ---------------------------------------------------------------------------

@pytest.fixture
def txt_file(tmp_path) -> Path:
    p = tmp_path / "sample.txt"
    p.write_text("Hello from a plain text document.\nSecond line here.", encoding="utf-8")
    return p


@pytest.fixture
def md_file(tmp_path) -> Path:
    p = tmp_path / "sample.md"
    p.write_text("# Heading\n\nThis is a markdown document with some content.", encoding="utf-8")
    return p


@pytest.fixture
def docx_file(tmp_path) -> Path:
    from docx import Document as DocxDocument
    doc = DocxDocument()
    doc.add_paragraph("Hello from a DOCX document.")
    doc.add_paragraph("This is the second paragraph.")
    out = tmp_path / "sample.docx"
    doc.save(str(out))
    return out


@pytest.fixture
def searchable_pdf_file(tmp_path) -> Path:
    import fitz
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 72), "Hello from a searchable PDF document.")
    out = tmp_path / "searchable.pdf"
    doc.save(str(out))
    doc.close()
    return out


@pytest.fixture
def scanned_pdf_file(tmp_path) -> Path:
    """A PDF with no text layer — simulates a scanned document."""
    import fitz
    doc = fitz.open()
    doc.new_page()  # blank page, no text inserted
    out = tmp_path / "scanned.pdf"
    doc.save(str(out))
    doc.close()
    return out


# ---------------------------------------------------------------------------
# Unit tests — text_extraction_service.extract() directly
# ---------------------------------------------------------------------------

def test_extract_txt_success(txt_file):
    result = text_extraction_service.extract(txt_file)
    assert result.extraction_status == "success"
    assert "Hello from a plain text document" in result.text
    assert result.text_length > 0
    assert result.text_preview != ""


def test_extract_md_success(md_file):
    result = text_extraction_service.extract(md_file)
    assert result.extraction_status == "success"
    assert "markdown document" in result.text
    assert result.text_length > 0


def test_extract_docx_success(docx_file):
    result = text_extraction_service.extract(docx_file)
    assert result.extraction_status == "success"
    assert "Hello from a DOCX document" in result.text
    assert result.text_length > 0


def test_extract_searchable_pdf_success(searchable_pdf_file):
    result = text_extraction_service.extract(searchable_pdf_file)
    assert result.extraction_status == "success"
    assert "Hello from a searchable PDF" in result.text
    assert result.text_length > 0


def test_extract_scanned_pdf_returns_failed(scanned_pdf_file):
    result = text_extraction_service.extract(scanned_pdf_file)
    assert result.extraction_status == "failed"
    assert result.text_length == 0
    assert result.text_preview == ""
    assert "scanned or image-based" in result.message


def test_extract_scanned_pdf_message_mentions_ocr(scanned_pdf_file):
    result = text_extraction_service.extract(scanned_pdf_file)
    assert "OCR" in result.message


def test_extract_preview_is_capped(tmp_path):
    p = tmp_path / "long.txt"
    p.write_text("A" * 1000, encoding="utf-8")
    result = text_extraction_service.extract(p)
    assert len(result.text_preview) <= text_extraction_service.PREVIEW_LENGTH


def test_extract_text_length_matches(txt_file):
    result = text_extraction_service.extract(txt_file)
    assert result.text_length == len(result.text)


# ---------------------------------------------------------------------------
# Integration tests — POST /api/documents/upload with real file content
# ---------------------------------------------------------------------------

@pytest.fixture(autouse=True)
def use_tmp_upload_dir(tmp_path):
    with patch.object(storage_service, "UPLOAD_DIR", tmp_path):
        yield


def _upload_bytes(filename: str, content: bytes, content_type: str):
    return client.post(
        "/api/documents/upload",
        files={"file": (filename, io.BytesIO(content), content_type)},
    )


def test_endpoint_txt_extraction_success(txt_file):
    response = _upload_bytes("notes.txt", txt_file.read_bytes(), "text/plain")
    data = response.json()
    assert response.status_code == 200
    assert data["extraction_status"] == "success"
    assert data["text_length"] > 0
    assert "Hello from a plain text document" in data["text_preview"]


def test_endpoint_md_extraction_success(md_file):
    response = _upload_bytes("readme.md", md_file.read_bytes(), "text/markdown")
    data = response.json()
    assert response.status_code == 200
    assert data["extraction_status"] == "success"
    assert data["text_length"] > 0


def test_endpoint_docx_extraction_success(docx_file):
    response = _upload_bytes(
        "contract.docx",
        docx_file.read_bytes(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    data = response.json()
    assert response.status_code == 200
    assert data["extraction_status"] == "success"
    assert data["text_length"] > 0


def test_endpoint_searchable_pdf_extraction_success(searchable_pdf_file):
    response = _upload_bytes("report.pdf", searchable_pdf_file.read_bytes(), "application/pdf")
    data = response.json()
    assert response.status_code == 200
    assert data["extraction_status"] == "success"
    assert data["text_length"] > 0


def test_endpoint_scanned_pdf_returns_failed_extraction(scanned_pdf_file):
    response = _upload_bytes("scan.pdf", scanned_pdf_file.read_bytes(), "application/pdf")
    data = response.json()
    assert response.status_code == 200  # file was uploaded; only extraction failed
    assert data["status"] == "uploaded"
    assert data["extraction_status"] == "failed"
    assert data["text_length"] == 0
    assert "scanned or image-based" in data["message"]
